#!/usr/bin/env python3

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "white-paper.md"
OUTPUT_DIR = ROOT / "output" / "docx"
OUTPUT = OUTPUT_DIR / "infiltration-and-inflow-technical-paper.docx"
FIGURE_PNG = ROOT / "generated" / "figure-png"

NAVY = "12233F"
BLUE = "176B87"
GRAY = "667085"
LIGHT = "E9EEF3"
PAPER = "F7F5F0"
ORANGE = "EF8354"
WHITE = "FFFFFF"
USABLE_DXA = 9360
TABLE_INDENT_DXA = 120


def set_font(run, name: str, size: float | None = None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)


def new_numbering_instance(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    style_num_id = doc.styles["List Number"].element.pPr.numPr.numId.val
    base_num = numbering.find(f"./w:num[@w:numId='{style_num_id}']", namespaces=numbering.nsmap)
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    used = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_id = max(used, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def add_inline_markup(paragraph, text: str, font_size: float = 10.5) -> None:
    text = re.sub(
        r"\\\((.*?)\\\)",
        lambda match: clean_math(match.group(1)).replace("\n", " "),
        text,
    )
    text = text.replace("\\&", "&").replace("\\_", "_")
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`|\*.*?\*|<https?://[^>]+>)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_font(run, "Calibri", font_size, color=NAVY)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, "Calibri", font_size, bold=True, color=NAVY)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, "Consolas", font_size - 0.5, color=BLUE)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, "Calibri", font_size, italic=True, color=NAVY)
        else:
            run = paragraph.add_run(token[1:-1])
            set_font(run, "Calibri", font_size, color=BLUE)
            run.underline = True
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_font(run, "Calibri", font_size, color=NAVY)


def clean_math(text: str) -> str:
    value = text.strip()
    value = value.replace("\\\\", "\n")
    value = value.replace("{,}", ",")
    value = value.replace("\\ ", " ")
    value = value.replace("\\begin{aligned}", "").replace("\\end{aligned}", "")
    value = value.replace("\\begin{cases}", "").replace("\\end{cases}", "")
    value = value.replace("\\left", "").replace("\\right", "")
    value = value.replace("\\quad", "  ").replace("\\,", " ")
    value = value.replace("\\times", " × ").replace("\\approx", " ≈ ")
    value = value.replace("\\leq", " ≤ ").replace("\\geq", " ≥ ")
    value = value.replace("\\mathchar`-", "-")
    value = value.replace("\\%", "%").replace("\\$", "$")

    # Remove text-oriented LaTeX wrappers before parsing operators. This also
    # eliminates braces nested inside fraction operands.
    for _ in range(12):
        replaced = re.sub(
            r"\\(?:mathrm|text|mathbf)\{([^{}]*)\}",
            r"\1",
            value,
        )
        if replaced == value:
            break
        value = replaced

    # Make grouped indices ordinary parenthetical text so nested fractions can
    # be reduced deterministically without exposing LaTeX commands to readers.
    for _ in range(12):
        replaced = re.sub(r"([_^])\{([^{}]*)\}", r"\1(\2)", value)
        if replaced == value:
            break
        value = replaced

    for _ in range(24):
        replaced = re.sub(
            r"\\frac\{([^{}]*)\}\s*\{([^{}]*)\}",
            r"(\1) / (\2)",
            value,
        )
        if replaced == value:
            break
        value = replaced

    value = re.sub(r"\\sqrt\{([^{}]*)\}", r"sqrt(\1)", value)
    value = re.sub(r"\\sqrt\s*([A-Za-z0-9_()^]+)", r"sqrt(\1)", value)
    value = re.sub(r"([A-Za-z0-9_)])sqrt\(", r"\1 sqrt(", value)
    value = re.sub(r"\\overline\{([^{}]*)\}", r"average(\1)", value)
    value = value.replace("\\sum", "SUM").replace("\\int", "INTEGRAL")
    value = value.replace("\\log", "log")
    value = value.replace("\\pi", "π")
    value = value.replace("\\Delta", "Δ").replace("\\Sigma", "Σ")
    value = value.replace("\\tau", "τ").replace("\\nu", "ν")
    value = value.replace("\\eta", "η").replace("\\epsilon", "ε")
    value = value.replace("\\infty", "∞")
    value = value.replace("&", " ")
    value = value.replace("{", "(").replace("}", ")")
    value = re.sub(r"\\([A-Za-z]+)", r"\1", value)
    value = re.sub(r"_\(([A-Za-z0-9]+)\)", r"_\1", value)
    value = re.sub(r"\^\(([A-Za-z0-9.+-]+)\)", r"^\1", value)
    value = re.sub(r"[ \t]+", " ", value)
    return "\n".join(line.strip() for line in value.splitlines() if line.strip())


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10

    if "Equation Block" not in doc.styles:
        style = doc.styles.add_style("Equation Block", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Cambria Math"
        style.font.size = Pt(9.5)
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.right_indent = Inches(0.25)
        style.paragraph_format.space_before = Pt(5)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.line_spacing = 1.05

    if "Code Block" not in doc.styles:
        style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style.font.size = Pt(8.5)
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.left_indent = Inches(0.2)
        style.paragraph_format.right_indent = Inches(0.2)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.line_spacing = 1.0


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)


def add_running_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("INFILTRATION AND INFLOW  |  TECHNICAL PAPER")
    set_font(run, "Calibri", 8.5, bold=True, color=GRAY)
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "D6DDE4")
    border.append(bottom)
    p_pr.append(border)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("Candidate engineering research  |  Page ")
    set_font(run, "Calibri", 8, color=GRAY)
    add_page_field(p)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TECHNICAL RESEARCH PAPER")
    set_font(r, "Calibri", 10, bold=True, color=ORANGE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Infiltration and Inflow")
    set_font(r, "Calibri", 30, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run(
        "A national engineering calculation framework with a "
        "Miami-Dade basin and pump-station case"
    )
    set_font(r, "Calibri", 15, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(80)
    r = p.add_run("Research paper and executable computational specification")
    set_font(r, "Calibri", 10.5, italic=True, color=GRAY)

    metadata = [
        ("VERSION", "1.0 candidate technical paper"),
        ("DATE", "July 27, 2026"),
        ("SAMPLE", "MD-EX-01 synthetic basin"),
        ("REGISTRY", "formula-register.yaml version 0.2.0"),
        ("BOUNDARY", "United States national layer with separate Miami-Dade rule pack"),
    ]
    table = doc.add_table(rows=len(metadata) + 1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "DOCUMENT CONTROL"
    table.rows[0].cells[1].text = "VALUE"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, BLUE)
        for run in cell.paragraphs[0].runs:
            set_font(run, "Calibri", 9, bold=True, color=WHITE)
    repeat_table_header(table.rows[0])
    for row_index, (label, value) in enumerate(metadata, start=1):
        left, right = table.rows[row_index].cells
        left.text = label
        right.text = value
        set_cell_shading(left, LIGHT)
        for run in left.paragraphs[0].runs:
            set_font(run, "Calibri", 9, bold=True, color=NAVY)
        for run in right.paragraphs[0].runs:
            set_font(run, "Calibri", 9.5, color=NAVY)
    set_table_geometry(table, [1700, 7660])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    r = p.add_run(
        "Synthetic engineering example. Not facility data, operating direction, "
        "capacity certification, or legal advice."
    )
    set_font(r, "Calibri", 9, bold=True, color=GRAY)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_figure(doc: Document, alt: str, source: str) -> None:
    png = FIGURE_PNG / Path(source).with_suffix(".png").name
    if not png.exists():
        raise FileNotFoundError(png)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    inline = run.add_picture(str(png), width=Inches(6.35))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt)
    p.paragraph_format.keep_with_next = True
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    r = caption.add_run(alt)
    set_font(r, "Calibri", 9, italic=True, color=GRAY)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    header = rows[0]
    body = rows[2:] if all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]) else rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for index, text in enumerate(header):
        cell = table.rows[0].cells[index]
        cell.text = ""
        p = cell.paragraphs[0]
        add_inline_markup(p, text.strip(), 9)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
        set_cell_shading(cell, BLUE)
    repeat_table_header(table.rows[0])
    for row_index, values in enumerate(body):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            p = cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if re.fullmatch(r"[\s$0-9,.\-%()]+", value) else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_markup(p, value.strip(), 8.8)
            if row_index % 2 == 1:
                set_cell_shading(cells[index], "F4F7F9")
    weights = []
    for index in range(len(header)):
        max_len = max(len(row[index]) if index < len(row) else 0 for row in [header] + body)
        weights.append(max(8, min(max_len, 42)))
    total_weight = sum(weights)
    widths = [round(USABLE_DXA * weight / total_weight) for weight in weights]
    widths[-1] += USABLE_DXA - sum(widths)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def parse_markdown(doc: Document, text: str) -> None:
    lines = text.splitlines()
    start = next(index for index, line in enumerate(lines) if line.startswith("## Document status"))
    index = start
    paragraph_buffer: list[str] = []
    active_numbering_id: int | None = None

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        value = " ".join(part.strip() for part in paragraph_buffer).strip()
        if value:
            p = doc.add_paragraph()
            add_inline_markup(p, value)
        paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            index += 1
            continue

        image_match = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", stripped)
        if image_match:
            flush_paragraph()
            add_figure(doc, image_match.group(1), image_match.group(2))
            index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            language = stripped[3:].strip()
            code: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index])
                index += 1
            p = doc.add_paragraph(style="Code Block")
            p.add_run("\n".join(code))
            set_paragraph_shading(p, LIGHT)
            index += 1
            continue

        if stripped == "\\[":
            flush_paragraph()
            equation: list[str] = []
            index += 1
            while index < len(lines) and lines[index].strip() != "\\]":
                equation.append(lines[index])
                index += 1
            p = doc.add_paragraph(style="Equation Block")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(clean_math("\n".join(equation)))
            set_paragraph_shading(p, "F1F5F7")
            index += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            table_rows: list[list[str]] = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                table_rows.append([cell.strip() for cell in candidate[1:-1].split("|")])
                index += 1
            add_markdown_table(doc, table_rows)
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = len(heading.group(1)) - 1
            if level == 1 and heading.group(2) in (
                "6. Universal calculation library",
                "7. Complete worked basin and pump-station example",
                "8. Agentic calculation architecture",
                "11. Formula wiring matrix",
                "12. Acronyms",
                "14. References",
            ):
                doc.add_page_break()
            doc.add_heading(heading.group(2), level=level)
            index += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or numbered:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet" if bullet else "List Number")
            add_inline_markup(p, (bullet or numbered).group(1))
            if numbered:
                if active_numbering_id is None:
                    active_numbering_id = new_numbering_instance(doc)
                apply_numbering(p, active_numbering_id)
            else:
                active_numbering_id = None
            index += 1
            continue

        active_numbering_id = None

        if stripped.startswith(">"):
            flush_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            add_inline_markup(p, stripped.lstrip("> ").strip())
            set_paragraph_shading(p, "FFF4EC")
            index += 1
            continue

        paragraph_buffer.append(stripped)
        index += 1
    flush_paragraph()


def build() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_running_header_footer(doc.sections[0])
    add_cover(doc)
    parse_markdown(doc, SOURCE.read_text())
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            set_keep_with_next(paragraph)
    doc.core_properties.title = "Infiltration and Inflow"
    doc.core_properties.subject = "National engineering calculation framework with a Miami-Dade basin and pump-station case"
    doc.core_properties.author = "OWOS technical research"
    doc.core_properties.keywords = "infiltration, inflow, RDII, RTK, pump station, Miami-Dade"
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    build()
